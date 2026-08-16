"""Build docs/PROJECT_REPORT.docx from docs/PROJECT_REPORT.md (print-friendly)."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "docs" / "PROJECT_REPORT.md"
OUT = ROOT / "docs" / "PROJECT_REPORT.docx"


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_formatted(paragraph, text: str, size=11):
    """Split **bold** segments."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def shade_header_row(row):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.makeelement(
            qn("w:shd"),
            {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "1F4E79"},
        )
        tcPr.append(shd)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.bold = True


def add_table(doc, rows: list[list[str]]):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val.strip())
            set_run_font(run, size=9, bold=(i == 0))
    shade_header_row(table.rows[0])
    doc.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        if re.match(r"^\|[\s:|-]+\|$", raw.replace(" ", "") if False else raw):
            cells = [c.strip() for c in raw.strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) or set(c) <= set("-: ") for c in cells):
                i += 1
                continue
        cells = [c.strip() for c in raw.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def build() -> Path:
    text = MD.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("CSE 471  ·  Parallel Log Intelligence  ·  Project Report")
    set_run_font(fr, size=9, color=RGBColor(0x55, 0x55, 0x55))

    styles = doc.styles["Normal"]
    styles.font.name = "Calibri"
    styles.font.size = Pt(11)
    pf = styles.paragraph_format
    pf.space_after = Pt(10)
    pf.line_spacing = 1.3

    i = 0
    in_code = False
    code_buf: list[str] = []
    title_done = False

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                set_run_font(run, name="Consolas", size=9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, i = parse_table(lines, i)
            if rows:
                width = max(len(r) for r in rows)
                norm = [r + [""] * (width - len(r)) for r in rows]
                add_table(doc, norm)
            continue

        if line.startswith("# ") and not title_done:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("PROJECT REPORT")
            set_run_font(run, size=14, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p2.add_run(line[2:].strip())
            set_run_font(run, size=22, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
            title_done = True
            i += 1
            continue

        if line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=1)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            i += 1
            continue
        if line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=2)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            i += 1
            continue

        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            add_formatted(p, line[2:].strip(), size=11)
            for r in p.runs:
                r.italic = True
            i += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted(p, line[2:].strip())
            i += 1
            continue

        if re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            add_formatted(p, re.sub(r"^\d+\.\s", "", line).strip())
            i += 1
            continue

        if line.startswith("**") and line.endswith("**") and line.count("**") == 2:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip("*").strip())
            set_run_font(run, size=12, bold=True)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        p = doc.add_paragraph()
        add_formatted(p, line.strip())
        i += 1

    try:
        doc.save(OUT)
        return OUT
    except PermissionError:
        alt = ROOT / "docs" / "PROJECT_REPORT_faculty.docx"
        doc.save(alt)
        return alt


if __name__ == "__main__":
    print(build())
